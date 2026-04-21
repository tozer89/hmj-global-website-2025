#!/usr/bin/env python3
"""Generate a varied 10-CV parser audit pack for HMJ sourcing trials."""

from __future__ import annotations

import json
import os
import subprocess
import sys
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    Image as ReportlabImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfgen import canvas


@dataclass
class CvSpec:
    id: str
    file_name: str
    format: str
    structure_tags: list[str]
    candidate_name: str
    current_title: str
    expected_parse_phrases: list[str]
    expected_shortlist_recommendation: str
    expected_parser: str
    expected_text_source: str
    requires_ocr: bool = False
    ocr_text: str = ""


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def build_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="HmJName",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        spaceAfter=10,
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        name="HmJTitle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#113355"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="HmJMeta",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#444444"),
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="HmJBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="HmJSection",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        textColor=colors.HexColor("#1F4B7A"),
        spaceBefore=8,
        spaceAfter=6,
    ))
    return styles


STYLES = build_styles()


def story_header(name: str, title: str, location: str, email: str, phone: str) -> list:
    return [
        Paragraph(name, STYLES["HmJName"]),
        Paragraph(title, STYLES["HmJTitle"]),
        Paragraph(f"{location} | {email} | {phone}", STYLES["HmJMeta"]),
        Spacer(1, 0.12 * inch),
    ]


def add_paragraphs(story: list, heading: str, paragraphs: Iterable[str]) -> None:
    story.append(Paragraph(heading, STYLES["HmJSection"]))
    for paragraph in paragraphs:
        story.append(Paragraph(paragraph, STYLES["HmJBody"]))


def write_pdf(path: Path, story: list) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=A4, topMargin=36, bottomMargin=36, leftMargin=42, rightMargin=42)
    doc.build(story)


def write_scanned_pdf(path: Path, image_path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    c.drawImage(ImageReader(str(image_path)), 24, 24, width=width - 48, height=height - 48, preserveAspectRatio=True, mask='auto')
    c.showPage()
    c.save()


def create_badge_image(path: Path, headline: str, subline: str) -> None:
    image = PILImage.new("RGB", (980, 320), color=(243, 248, 252))
    draw = ImageDraw.Draw(image)
    font_large = ImageFont.load_default()
    font_small = ImageFont.load_default()
    draw.rounded_rectangle((25, 25, 955, 295), radius=28, outline=(34, 78, 122), width=4, fill=(233, 242, 249))
    draw.text((70, 90), headline, fill=(14, 42, 66), font=font_large)
    draw.text((70, 150), subline, fill=(52, 72, 92), font=font_small)
    draw.text((70, 205), "Commissioning | QA/QC | Handover | Data Centre", fill=(52, 72, 92), font=font_small)
    image.save(path)


def create_scanned_cv_image(path: Path, lines: list[str]) -> None:
    image = PILImage.new("L", (1600, 2200), color=255)
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    y = 70
    for line in lines:
        draw.text((80, y), line, fill=18, font=font)
        y += 42
        if line.endswith(":"):
            y += 8
    image.save(path)


def write_docx(path: Path, name: str, title: str, location: str, email: str, phone: str, summary: list[str], experience: list[dict], skills: list[str], image_path: Path | None = None, table_rows: list[list[str]] | None = None) -> None:
    document = Document()
    document.add_heading(name, 0)
    document.add_paragraph(title)
    document.add_paragraph(f"{location} | {email} | {phone}")

    if image_path is not None and image_path.exists():
        document.add_picture(str(image_path), width=Inches(2.2))

    document.add_heading("Professional Summary", level=1)
    for paragraph in summary:
        document.add_paragraph(paragraph)

    document.add_heading("Key Skills", level=1)
    document.add_paragraph(", ".join(skills))

    if table_rows:
        document.add_heading("Selected Project Portfolio", level=1)
        table = document.add_table(rows=1, cols=len(table_rows[0]))
        table.style = "Table Grid"
        headers = table.rows[0].cells
        for idx, header in enumerate(table_rows[0]):
            headers[idx].text = header
        for row in table_rows[1:]:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = value

    document.add_heading("Employment History", level=1)
    for role in experience:
        document.add_paragraph(f"{role['dates']} | {role['title']} | {role['employer']}")
        for bullet in role["bullets"]:
            document.add_paragraph(bullet, style="List Bullet")

    document.save(path)


def write_legacy_doc(path: Path, name: str, title: str, location: str, email: str, paragraphs: list[str], experience: list[dict]) -> None:
    rtf_lines = [
        r"{\rtf1\ansi\deff0",
        r"{\fonttbl{\f0 Helvetica;}}",
        r"\f0\fs28 " + f"{name}\\par",
        r"\fs24 " + f"{title}\\par",
        f"{location} | {email}\\par",
        r"\fs22 Profile\\par",
    ]
    for paragraph in paragraphs:
        rtf_lines.append(paragraph.replace("\\", "\\\\").replace("{", r"\{").replace("}", r"\}") + r"\par")
    rtf_lines.append(r"\fs22 Experience\\par")
    for role in experience:
        rtf_lines.append(f"{role['dates']} | {role['title']} | {role['employer']}\\par")
        for bullet in role["bullets"]:
            rtf_lines.append(f"- {bullet}\\par")
    rtf_lines.append("}")

    rtf_path = path.with_suffix(".rtf")
    rtf_path.write_text("\n".join(rtf_lines), encoding="utf-8")
    subprocess.run(
        ["/usr/bin/textutil", "-convert", "doc", "-output", str(path), str(rtf_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if rtf_path.exists():
        rtf_path.unlink()


def make_pdf_candidate_one(output_dir: Path) -> CvSpec:
    file_name = "amelia-hart-senior-electrical-commissioning-manager.pdf"
    story = story_header(
        "Amelia Hart",
        "Senior Electrical Commissioning Manager",
        "Leeds",
        "amelia.hart@example.com",
        "07700 900101",
    )
    add_paragraphs(story, "Professional Summary", [
        "Senior Electrical Commissioning Manager with 14 years across hyperscale data centres, mission critical commissioning, package leadership, QA/QC close-out and integrated systems handover on live and fast-track builds.",
        "Strong record of coordinating subcontractors, witnessing testing, driving snag close-out and taking projects from energisation planning through IST and final client handover.",
    ])
    add_paragraphs(story, "Core Skills", [
        "Electrical commissioning, QA/QC, commissioning scripts, subcontractor management, package management, integrated systems testing, permit control, client reporting.",
    ])
    add_paragraphs(story, "Career History", [
        "2023 - Present | Senior Electrical Commissioning Manager | NorthBridge Mission Critical. Led electrical commissioning and handover on a 54MW data centre campus in West Yorkshire, coordinating Tier 1 subcontractors, QA/QC and IST readiness.",
        "2020 - 2023 | Electrical Commissioning Manager | Arcus Build. Managed commissioning scripts, witnessing, red-line resolution and final electrical package sign-off for a regulated pharma expansion in Leeds.",
        "2016 - 2020 | Electrical Package Manager | Integrum Projects. Delivered MV/LV packages and client-facing handover packs across cleanroom and mission critical environments.",
    ])
    write_pdf(output_dir / file_name, story)
    return CvSpec(
        id="cv01",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "multi_page", "direct_match", "mission_critical"],
        candidate_name="Amelia Hart",
        current_title="Senior Electrical Commissioning Manager",
        expected_parse_phrases=["Amelia Hart", "Senior Electrical Commissioning Manager", "mission critical commissioning", "QA/QC", "subcontractors"],
        expected_shortlist_recommendation="strong",
        expected_parser="pdf-parse",
        expected_text_source="native_pdf_text",
    )


def make_docx_candidate_two(output_dir: Path) -> CvSpec:
    file_name = "daniel-osei-electrical-project-manager.docx"
    write_docx(
        output_dir / file_name,
        "Daniel Osei",
        "Electrical Project Manager",
        "Manchester",
        "daniel.osei@example.com",
        "07700 900102",
        [
            "Electrical Project Manager with direct leadership of data centre and pharma packages, including subcontractor management, QA/QC sign-off, commissioning planning and client reporting across the North.",
            "Comfortable bridging delivery and commissioning teams to drive programme, energisation readiness and handover quality.",
        ],
        [
            {
                "dates": "2022 - Present",
                "title": "Electrical Project Manager",
                "employer": "Vector Mission Critical",
                "bullets": [
                    "Led electrical package delivery on a 36MW data centre fit-out with subcontractor coordination and weekly QA/QC reviews.",
                    "Managed commissioning readiness, witness testing and handover packs for mission critical switchgear and UPS packages.",
                ],
            },
            {
                "dates": "2018 - 2022",
                "title": "Electrical Package Manager",
                "employer": "Helix Projects",
                "bullets": [
                    "Drove cleanroom electrical packages for regulated pharma clients in Yorkshire and the Midlands.",
                ],
            },
        ],
        ["Electrical package management", "QA/QC", "Subcontractor coordination", "Commissioning planning", "Data centre delivery", "SMSTS"],
    )
    return CvSpec(
        id="cv02",
        file_name=file_name,
        format="docx",
        structure_tags=["docx", "standard", "direct_match"],
        candidate_name="Daniel Osei",
        current_title="Electrical Project Manager",
        expected_parse_phrases=["Daniel Osei", "Electrical Project Manager", "data centre", "QA/QC", "commissioning readiness"],
        expected_shortlist_recommendation="strong",
        expected_parser="mammoth",
        expected_text_source="docx_text",
    )


def make_docx_candidate_three(output_dir: Path) -> CvSpec:
    file_name = "marek-novak-electrical-package-manager-table.docx"
    write_docx(
        output_dir / file_name,
        "Marek Novak",
        "Electrical Package Manager",
        "Bradford",
        "marek.novak@example.com",
        "07700 900103",
        [
            "Electrical Package Manager with 12 years of mission critical and cleanroom delivery across MCC, UPS, containment, testing and handover.",
            "Known for strong package control, trade sequencing, punch list ownership and structured close-out reporting.",
        ],
        [
            {
                "dates": "2021 - Present",
                "title": "Electrical Package Manager",
                "employer": "Catalyst Critical Systems",
                "bullets": [
                    "Owned package planning, subcontractor management and QA/QC close-out on hyperscale data centre workstreams.",
                    "Worked with commissioning managers to align energisation, IST and client handover milestones.",
                ],
            },
            {
                "dates": "2017 - 2021",
                "title": "Electrical Site Manager",
                "employer": "PureClean Engineering",
                "bullets": [
                    "Delivered cleanroom and pharma packages with strict permit, quality and test documentation standards.",
                ],
            },
        ],
        ["Package management", "Subcontractor coordination", "QA/QC", "Data centres", "Cleanrooms", "ECS", "SSSTS"],
        table_rows=[
            ["Dates", "Client / Site", "Package / Scope", "Sector / Outcome"],
            ["2024 - 2025", "Leeds DC-3", "MCC, UPS, busbar, QA/QC", "Data centre / handover achieved"],
            ["2022 - 2024", "West Yorkshire cleanroom", "Containment, power, testing", "Cleanroom / zero major defects"],
            ["2020 - 2022", "North pharma expansion", "Electrical package and commissioning interface", "Pharma / GMP close-out"],
        ],
    )
    return CvSpec(
        id="cv03",
        file_name=file_name,
        format="docx",
        structure_tags=["docx", "table_heavy", "direct_match"],
        candidate_name="Marek Novak",
        current_title="Electrical Package Manager",
        expected_parse_phrases=["Marek Novak", "Electrical Package Manager", "hyperscale data centre", "cleanroom", "QA/QC"],
        expected_shortlist_recommendation="strong",
        expected_parser="mammoth",
        expected_text_source="docx_text",
    )


def make_pdf_candidate_four(output_dir: Path) -> CvSpec:
    file_name = "priya-nair-commissioning-lead-long-form.pdf"
    story = story_header(
        "Priya Nair",
        "Commissioning Lead",
        "Birmingham",
        "priya.nair@example.com",
        "07700 900104",
    )
    add_paragraphs(story, "Executive Profile", [
        "Commissioning Lead with 15 years in electrical delivery, mission critical coordination, commissioning management, integrated testing and highly controlled client handovers across EMEA data centre and life sciences programmes.",
        "Known for dense but rigorous documentation, structured issue burn-down, FAT/SAT governance and practical site leadership in live environments.",
    ])
    for page_number in range(1, 10):
        add_paragraphs(story, f"Programme {page_number}", [
            "2024 - Present | Commissioning Lead | Helion Critical Infrastructure. Directed commissioning interfaces, subcontractor sequencing, QA/QC evidence control, test packs, live risk review and multi-party sign-off on a complex mission critical campus. Produced detailed narratives covering isolations, energisation sequencing, witness points, defect burn-down, red-line control, commissioning scripts, Level 4/5 programme ties and integrated systems dependencies.",
            "2019 - 2024 | Electrical Commissioning Manager | Axis Delivery. Supported cleanroom, pharma and advanced manufacturing environments where exacting documentation, punch resolution and careful handover evidence were critical to successful close-out.",
        ])
        story.append(PageBreak())
    add_paragraphs(story, "Credentials", [
        "SMSTS, ECS, AP awareness, commissioning scripts, QA/QC leadership, witness testing, subcontractor management, handover reporting.",
    ])
    write_pdf(output_dir / file_name, story)
    return CvSpec(
        id="cv04",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "ten_page", "dense_wording", "adjacent_match"],
        candidate_name="Priya Nair",
        current_title="Commissioning Lead",
        expected_parse_phrases=["Priya Nair", "Commissioning Lead", "mission critical", "integrated systems", "QA/QC"],
        expected_shortlist_recommendation="strong",
        expected_parser="pdf-parse",
        expected_text_source="native_pdf_text",
    )


def make_pdf_candidate_five(output_dir: Path) -> CvSpec:
    file_name = "chloe-bennett-electrical-construction-manager-basic.pdf"
    story = story_header(
        "Chloe Bennett",
        "Electrical Construction Manager",
        "Leeds",
        "chloe.bennett@example.com",
        "07700 900105",
    )
    add_paragraphs(story, "Summary", [
        "Electrical Construction Manager with site delivery, subcontractor coordination and quality close-out experience on commercial and mission critical projects.",
        "Ready to step back into a package-led role with more direct commissioning exposure.",
    ])
    add_paragraphs(story, "Experience", [
        "2022 - Present | Electrical Construction Manager | Northern Build. Managed electrical trade sequencing, subcontractors and QA snag close-out on a mission critical fit-out in Yorkshire.",
        "2018 - 2022 | Senior Electrical Supervisor | Midtown Projects. Oversaw installation quality, programme tracking and trade interfaces on large commercial and light industrial schemes.",
    ])
    write_pdf(output_dir / file_name, story)
    return CvSpec(
        id="cv05",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "one_page", "basic", "adjacent_match"],
        candidate_name="Chloe Bennett",
        current_title="Electrical Construction Manager",
        expected_parse_phrases=["Chloe Bennett", "Electrical Construction Manager", "subcontractor coordination", "mission critical"],
        expected_shortlist_recommendation="possible",
        expected_parser="pdf-parse",
        expected_text_source="native_pdf_text",
    )


def make_docx_candidate_six(output_dir: Path, image_path: Path) -> CvSpec:
    file_name = "tomasz-zielinski-senior-commissioning-engineer-image.docx"
    write_docx(
        output_dir / file_name,
        "Tomasz Zielinski",
        "Senior Commissioning Engineer",
        "Sheffield",
        "tomasz.zielinski@example.com",
        "07700 900106",
        [
            "Senior Commissioning Engineer with electrical commissioning, witness testing, snag close-out and practical handover experience across data centre and advanced manufacturing work.",
            "Acts as a strong site interface between electrical delivery, controls vendors and client representatives.",
        ],
        [
            {
                "dates": "2023 - Present",
                "title": "Senior Commissioning Engineer",
                "employer": "Vertex Commissioning",
                "bullets": [
                    "Supported mission critical energisation, test witnessing and close-out documentation.",
                    "Worked closely with package managers and subcontractors to drive defects to completion.",
                ],
            },
            {
                "dates": "2019 - 2023",
                "title": "Electrical Commissioning Engineer",
                "employer": "Northern Controls & Power",
                "bullets": [
                    "Delivered testing and handover support on data centre and pharmaceutical upgrades.",
                ],
            },
        ],
        ["Electrical commissioning", "Witness testing", "Subcontractor coordination", "Data centres", "Pharma", "ECS"],
        image_path=image_path,
    )
    return CvSpec(
        id="cv06",
        file_name=file_name,
        format="docx",
        structure_tags=["docx", "embedded_image", "adjacent_match"],
        candidate_name="Tomasz Zielinski",
        current_title="Senior Commissioning Engineer",
        expected_parse_phrases=["Tomasz Zielinski", "Senior Commissioning Engineer", "mission critical energisation", "pharmaceutical"],
        expected_shortlist_recommendation="possible",
        expected_parser="mammoth",
        expected_text_source="docx_text",
    )


def make_legacy_doc_candidate_seven(output_dir: Path) -> CvSpec:
    file_name = "patrick-reilly-electrical-site-manager-legacy.doc"
    write_legacy_doc(
        output_dir / file_name,
        "Patrick Reilly",
        "Electrical Site Manager",
        "Leeds",
        "patrick.reilly@example.com",
        [
            "Electrical Site Manager with mission critical and cleanroom projects, subcontractor management, QA/QC inspections and structured handover packs.",
            "Experience coordinating electrical packages, commissioning activities and snag resolution across live and fast-track environments.",
        ],
        [
            {
                "dates": "2021 - Present",
                "title": "Electrical Site Manager",
                "employer": "NorthWest Site Services",
                "bullets": [
                    "Managed electrical package delivery and subcontractors on a data centre expansion in Yorkshire.",
                    "Worked with commissioning teams on QA/QC close-out and witness activities.",
                ],
            },
            {
                "dates": "2017 - 2021",
                "title": "Electrical Supervisor",
                "employer": "PureClean Installations",
                "bullets": [
                    "Delivered cleanroom and pharma electrical projects with strong documentation discipline.",
                ],
            },
        ],
    )
    return CvSpec(
        id="cv07",
        file_name=file_name,
        format="doc",
        structure_tags=["legacy_doc", "older_word_version", "adjacent_match"],
        candidate_name="Patrick Reilly",
        current_title="Electrical Site Manager",
        expected_parse_phrases=["Patrick Reilly", "Electrical Site Manager", "data centre expansion", "QA/QC"],
        expected_shortlist_recommendation="strong",
        expected_parser="textutil",
        expected_text_source="legacy_doc_textutil",
    )


def make_scanned_pdf_candidate_eight(output_dir: Path, scan_image_path: Path) -> CvSpec:
    file_name = "luca-romano-senior-electrical-supervisor-scan.pdf"
    ocr_text = "\n".join([
        "Luca Romano",
        "Senior Electrical Supervisor",
        "Leeds",
        "luca.romano@example.com",
        "Profile",
        "Senior Electrical Supervisor with data centre and pharma experience, subcontractor coordination, QA close-out and support for commissioning and handover.",
        "Experience",
        "2022 - Present | Senior Electrical Supervisor | Orbit Commissioning Support",
        "2018 - 2022 | Electrical Supervisor | North Pharma Projects",
        "Skills",
        "Subcontractor management, QA/QC, commissioning support, cleanroom delivery, mission critical fit-out",
    ])
    create_scanned_cv_image(scan_image_path, ocr_text.splitlines())
    write_scanned_pdf(output_dir / file_name, scan_image_path)
    return CvSpec(
        id="cv08",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "scanned_image_only", "ocr_required", "adjacent_match"],
        candidate_name="Luca Romano",
        current_title="Senior Electrical Supervisor",
        expected_parse_phrases=["Luca Romano", "Senior Electrical Supervisor", "data centre", "commissioning support"],
        expected_shortlist_recommendation="possible",
        expected_parser="pdf-ocr",
        expected_text_source="ocr_pdf_text",
        requires_ocr=True,
        ocr_text=ocr_text,
    )


def make_pdf_candidate_nine(output_dir: Path) -> CvSpec:
    file_name = "ben-carter-mechanical-commissioning-manager.pdf"
    story = story_header(
        "Ben Carter",
        "Mechanical Commissioning Manager",
        "Manchester",
        "ben.carter@example.com",
        "07700 900109",
    )
    add_paragraphs(story, "Profile", [
        "Mechanical Commissioning Manager with chilled water, HVAC, BMS and water treatment focus across hospitals and commercial towers.",
        "Strong mechanical trade focus, but no electrical package ownership or electrical commissioning leadership.",
    ])
    add_paragraphs(story, "Experience", [
        "2021 - Present | Mechanical Commissioning Manager | Thermex Projects. Led mechanical commissioning across HVAC, water systems and BMS integrations on healthcare and office schemes.",
        "2017 - 2021 | Mechanical Project Engineer | Urban Building Services. Managed mechanical testing, balancing and handover documentation.",
    ])
    write_pdf(output_dir / file_name, story)
    return CvSpec(
        id="cv09",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "wrong_discipline", "reject"],
        candidate_name="Ben Carter",
        current_title="Mechanical Commissioning Manager",
        expected_parse_phrases=["Ben Carter", "Mechanical Commissioning Manager", "HVAC"],
        expected_shortlist_recommendation="reject",
        expected_parser="pdf-parse",
        expected_text_source="native_pdf_text",
    )


def make_pdf_candidate_ten(output_dir: Path) -> CvSpec:
    file_name = "sophie-webb-facilities-maintenance-electrician-basic.pdf"
    story = story_header(
        "Sophie Webb",
        "Facilities Maintenance Electrician",
        "Wakefield",
        "sophie.webb@example.com",
        "07700 900110",
    )
    add_paragraphs(story, "Summary", [
        "Facilities Maintenance Electrician focused on planned maintenance, reactive calls, landlord compliance checks and small works across retail and office estates.",
    ])
    add_paragraphs(story, "Experience", [
        "2020 - Present | Facilities Maintenance Electrician | Civic Estates. Delivered maintenance and fault-finding across occupied commercial sites.",
        "2016 - 2020 | Electrician | Local Building Works. Carried out small power, lighting and testing work on refurbishments.",
    ])
    write_pdf(output_dir / file_name, story)
    return CvSpec(
        id="cv10",
        file_name=file_name,
        format="pdf",
        structure_tags=["pdf", "basic", "maintenance_only", "reject"],
        candidate_name="Sophie Webb",
        current_title="Facilities Maintenance Electrician",
        expected_parse_phrases=["Sophie Webb", "Facilities Maintenance Electrician", "reactive calls"],
        expected_shortlist_recommendation="reject",
        expected_parser="pdf-parse",
        expected_text_source="native_pdf_text",
    )


def build_trial_pack(output_dir: Path) -> list[dict]:
    ensure_dir(output_dir)
    assets_dir = ensure_dir(output_dir / "_assets")
    badge_path = assets_dir / "commissioning-badge.png"
    scan_image_path = assets_dir / "luca-romano-scan-source.png"
    create_badge_image(badge_path, "COMMISSIONING DELIVERY", "QA/QC close-out | Data centre package leadership")

    specs = [
        make_pdf_candidate_one(output_dir),
        make_docx_candidate_two(output_dir),
        make_docx_candidate_three(output_dir),
        make_pdf_candidate_four(output_dir),
        make_pdf_candidate_five(output_dir),
        make_docx_candidate_six(output_dir, badge_path),
        make_legacy_doc_candidate_seven(output_dir),
        make_scanned_pdf_candidate_eight(output_dir, scan_image_path),
        make_pdf_candidate_nine(output_dir),
        make_pdf_candidate_ten(output_dir),
    ]

    manifest = []
    for spec in specs:
        manifest.append({
            "id": spec.id,
            "file_name": spec.file_name,
            "format": spec.format,
            "structure_tags": spec.structure_tags,
            "candidate_name": spec.candidate_name,
            "current_title": spec.current_title,
            "expected_parse_phrases": spec.expected_parse_phrases,
            "expected_shortlist_recommendation": spec.expected_shortlist_recommendation,
            "expected_parser": spec.expected_parser,
            "expected_text_source": spec.expected_text_source,
            "requires_ocr": spec.requires_ocr,
            "ocr_text": spec.ocr_text,
            "relative_path": spec.file_name,
        })
    return manifest


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: generate_trial_cv_pack.py <output-dir>", file=sys.stderr)
        return 1
    output_dir = Path(argv[1]).expanduser().resolve()
    ensure_dir(output_dir)
    manifest = build_trial_pack(output_dir)
    (output_dir / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    print(f"Generated {len(manifest)} trial CV files in {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
